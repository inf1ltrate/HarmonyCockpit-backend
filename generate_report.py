from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_img_placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'【📷 此处插入{text}】')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('《计算机编程实践》')
run.font.size = Pt(26); run.font.bold = True
run.font.name = '黑体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('四级项目报告')
run.font.size = Pt(26); run.font.bold = True
run.font.name = '黑体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('（2025-2026学年 第3学期）')
run.font.size = Pt(16)
run.font.name = '宋体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

for label, value in [('项目名称', '鸿蒙智能座舱车主端后端系统'), ('项目内容', '车主端后端API实现'), ('学号', '（请填写）'), ('姓名', '（请填写）')]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)
    run.font.name = '宋体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============ 目录 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目  录')
run.font.size = Pt(18); run.font.bold = True
run.font.name = '黑体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for item, page in [('一、需求分析', '1'), ('二、数据库设计', '2'), ('三、接口设计', '3'), ('四、项目实现', '5'), ('五、亮点功能', '7'), ('六、项目总结', '8')]:
    p = doc.add_paragraph()
    run = p.add_run(f'{item}{"." * (28 - len(item))}{page}')
    run.font.size = Pt(14)
    run.font.name = '宋体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============ 一、需求分析 ============
add_heading_cn(doc, '一、需求分析', 1)

add_body(doc, '本项目为鸿蒙智能座舱车主端后端系统，基于 Node.js + Express + MySQL 技术栈，采用 MVC 分层架构，为鸿蒙智能座舱（车主端）提供 RESTful API 服务。系统面向车主用户，提供车辆信息查看、远程控制、意见反馈等核心功能。', indent=True)

add_body(doc, '车主端核心业务需求：', indent=True)

for i, item in enumerate([
    '登录认证：车主通过用户ID和密码登录系统，系统验证身份后返回个人信息。支持用户不存在、密码错误、参数缺失等异常场景的处理。',
    '查看和维护个人信息：车主登录后可查看自己的姓名、性别、出生日期等详细信息，并支持修改个人资料（姓名、性别、出生日期、密码）。',
    '查看和维护车辆信息：车主可查看名下所有车辆的详细信息，包括品牌、功率、最大功率、扭矩、电池容量、续航里程、油耗等参数。支持分页查询和按用户ID过滤，通过LEFT JOIN关联用户表获取车主姓名。',
    '远程车辆控制：支持对绑定车辆进行远程操控，包括：空调开关及温度调节、制冷模式切换、空气循环（内循环/外循环）、除雾、车门开关、灯光开关、风扇开关、车辆移动（前进/后退/左转/右转/停车）。所有操作记录保存到数据库，可通过操作记录页面查看历史行为。',
    '意见反馈：车主可提交使用过程中的反馈建议，查看所有用户提交的反馈信息列表，支持新增反馈功能。',
    '数据统计：展示最近7日车辆行驶里程折线图，统计本周总里程、日均里程和最长单日里程。',
], 1):
    add_body(doc, f'{i}. {item}', indent=True)

doc.add_page_break()

# ============ 二、数据库设计 ============
add_heading_cn(doc, '二、数据库设计', 1)

add_body(doc, '本系统使用 MySQL 8.0 数据库，数据库名为 car，字符集 utf8mb4。共设计 4 张核心表，2 个视图，2 个索引。', indent=True)

add_body(doc, '2.1 user 表（用户信息表）', indent=True)
add_body(doc, '用于存储车主用户信息，包含用户ID、姓名、性别、出生日期、密码等字段。', indent=True)
add_table(doc, ['字段名', '类型', '约束', '说明'], [
    ['user_id', 'CHAR(15)', 'PRIMARY KEY', '用户ID（业务主键）'],
    ['userName', 'VARCHAR(20)', 'NOT NULL', '姓名'],
    ['Sex', 'VARCHAR(2)', 'NOT NULL', '性别（男/女）'],
    ['Birthdate', 'DATE', 'NOT NULL', '出生日期'],
    ['password', 'VARCHAR(20)', 'NOT NULL', '登录密码'],
])
add_body(doc, '表1  user 表结构', indent=False)
doc.add_paragraph()

add_body(doc, '2.2 info 表（车辆信息表）', indent=True)
add_body(doc, '用于存储车辆详细信息，user_id 外键关联 user 表。', indent=True)
add_table(doc, ['字段名', '类型', '约束', '说明'], [
    ['car_id', 'CHAR(10)', 'PRIMARY KEY', '车辆ID'],
    ['carName', 'VARCHAR(20)', 'NOT NULL', '车辆名称'],
    ['brand', 'VARCHAR(10)', 'NOT NULL', '品牌'],
    ['frame_id', 'CHAR(50)', 'NOT NULL', '车架号(VIN)'],
    ['user_id', 'CHAR(15)', 'FK→user', '所属用户ID'],
    ['power', 'INT', 'CHECK>0', '功率(kW)'],
    ['max_power', 'INT', 'CHECK>0', '最大功率(kW)'],
    ['torque', 'INT', 'CHECK>0', '扭矩(N·m)'],
    ['battery', 'INT', 'CHECK>0', '电池容量(kWh)'],
    ['miles', 'INT', 'CHECK>0', '续航里程(km)'],
    ['mpg', 'FLOAT', 'CHECK>0', '油耗(L/100km)'],
])
add_body(doc, '表2  info 表结构', indent=False)
doc.add_paragraph()

add_body(doc, '2.3 state 表（车辆操控表）', indent=True)
add_body(doc, '存储车辆实时控制状态，复合主键 (user_ID, car_ID, time)，user_ID 和 car_ID 为外键。', indent=True)
add_table(doc, ['字段名', '类型', '默认值', '说明'], [
    ['user_ID', 'CHAR(15)', 'NOT NULL', '用户ID（FK→user）'],
    ['car_ID', 'CHAR(10)', 'NOT NULL', '车辆ID（FK→info）'],
    ['ac_on', 'CHAR(2)', "'0'", '空调开关(0关/1开)'],
    ['ac_temp', 'FLOAT', '-', '空调温度(°C)'],
    ['fan', 'CHAR(2)', "'0'", '风扇'],
    ['rec', 'CHAR(2)', "'0'", '内循环(0外/1内)'],
    ['fog', 'CHAR(2)', "'0'", '除雾'],
    ['light', 'CHAR(2)', "'0'", '灯光'],
    ['door', 'CHAR(2)', "'0'", '车门(0关/1开)'],
    ['stop1', 'CHAR(2)', "'0'", '停车'],
    ['left1', 'CHAR(2)', "'0'", '左转'],
    ['right1', 'CHAR(2)', "'0'", '右转'],
    ['move', 'CHAR(2)', "'0'", '前进'],
    ['back', 'CHAR(2)', "'0'", '后退'],
    ['time', 'DATETIME(3)', 'NOT NULL', '状态时间(毫秒精度)'],
])
add_body(doc, '表3  state 表结构（复合主键）', indent=False)
doc.add_paragraph()

add_body(doc, '2.4 feedback 表（意见反馈表）', indent=True)
add_body(doc, '存储车主提交的意见反馈，user_id 外键关联 user 表。', indent=True)
add_table(doc, ['字段名', '类型', '约束', '说明'], [
    ['idea_id', 'CHAR(10)', 'PRIMARY KEY', '反馈ID'],
    ['ideaName', 'VARCHAR(20)', 'NOT NULL', '反馈标题'],
    ['content', 'VARCHAR(100)', 'NOT NULL', '反馈内容'],
    ['time', 'DATETIME', 'NOT NULL', '提交时间'],
    ['user_id', 'CHAR(15)', 'FK→user', '反馈用户ID'],
])
add_body(doc, '表4  feedback 表结构', indent=False)
doc.add_paragraph()

add_body(doc, '2.5 视图', indent=True)
add_body(doc, 'xuqing_car：查询车主"许青"的车辆信息（车主昵称、汽车名、汽车品牌、汽车里程）。', indent=True)
add_body(doc, 'xuqing_carstate：查询车主"许青"的车辆操控状态（车主昵称、汽车名、空调开关、空调温度、左转、右转、操作时间）。', indent=True)

add_body(doc, '2.6 索引', indent=True)
add_body(doc, 'IDX_mpg：info 表的 mpg 字段索引。', indent=True)
add_body(doc, 'IDX_miles：info 表的 miles 字段索引。', indent=True)

add_body(doc, '2.7 初始数据', indent=True)
add_body(doc, '用户5条：王琦、许青、罗锋、陈平安、石昊。车辆11条：含小马驹(特斯拉)、海豹(比亚迪)、猎豹(问界)、M7(问界)、S7(智界)、S800(尊界)等。操控记录31条。反馈3条。', indent=True)

doc.add_page_break()

# ============ 三、接口设计 ============
add_heading_cn(doc, '三、接口设计', 1)

add_body(doc, '系统采用 RESTful API 设计风格，Base URL：http://localhost:3000/api，统一 JSON 响应格式：', indent=True)
add_code(doc, '''{
    "success": true,
    "data": { ... },
    "message": "操作成功"
}''')
add_code(doc, '''{
    "success": false,
    "error": { "code": "ERROR_CODE", "message": "错误描述" }
}''')

add_body(doc, 'HTTP 状态码：200(OK)、201(Created)、400(Bad Request)、401(Unauthorized)、404(Not Found)、500(Server Error)。', indent=True)

add_body(doc, '3.1 核心 API 接口列表', indent=True)
add_table(doc, ['模块', '方法', '路径', '说明'], [
    ['登录', 'POST', '/api/users/login', '用户登录认证'],
    ['用户', 'GET', '/api/users/:user_id', '查询用户详情'],
    ['用户', 'PUT', '/api/users/:user_id', '更新用户信息'],
    ['车辆', 'GET', '/api/infos', '分页查询车辆列表'],
    ['车辆', 'GET', '/api/infos/withUser', 'LEFT JOIN 用户名的车辆列表'],
    ['车辆', 'GET', '/api/infos/:car_id', '查询车辆详情'],
    ['车辆', 'POST', '/api/infos', '新增车辆'],
    ['车辆', 'PUT', '/api/infos/:car_id', '更新车辆信息'],
    ['车辆', 'DELETE', '/api/infos/:car_id', '删除车辆'],
    ['控制', 'GET', '/api/states', '分页查询状态记录'],
    ['控制', 'GET', '/api/states/latest', '查询最新状态'],
    ['控制', 'GET', '/api/states/stats/miles', '近7日里程统计'],
    ['控制', 'POST', '/api/states', '新增状态记录'],
    ['控制', 'PUT', '/api/states', '更新最新状态'],
    ['控制', 'DELETE', '/api/states', '删除状态记录'],
    ['反馈', 'GET', '/api/feedbacks', '分页查询反馈'],
    ['反馈', 'GET', '/api/feedbacks/:idea_id', '查询反馈详情'],
    ['反馈', 'POST', '/api/feedbacks', '新增反馈'],
    ['反馈', 'PUT', '/api/feedbacks/:idea_id', '更新反馈'],
    ['反馈', 'DELETE', '/api/feedbacks/:idea_id', '删除反馈'],
])

add_body(doc, '3.2 登录接口示例', indent=True)
add_body(doc, 'POST /api/users/login', indent=True)
add_code(doc, '请求体: { "user_id": "277273737937", "password": "12345678" }')
add_code(doc, '成功响应: { "success": true, "data": { "user_id": "277273737937", "userName": "王琦", "Sex": "男", "Birthdate": "2009-03-01" }, "message": "登录成功" }')
add_code(doc, '失败响应: { "success": false, "error": { "code": "UNAUTHORIZED", "message": "密码错误" } }')

add_body(doc, '3.3 车辆控制接口示例', indent=True)
add_body(doc, 'POST /api/states', indent=True)
add_code(doc, '''请求体: {
    "user_ID": "277273737937", "car_ID": "C0009",
    "ac_on": "1", "ac_temp": 25, "rec": "1", "fog": "0",
    "time": "2026-07-06 14:45:00.789"
}''')
add_code(doc, '成功响应: { "success": true, "message": "新增成功" }')

add_img_placeholder(doc, 'Postman 测试截图（至少3张：登录成功、登录失败、车辆控制操作）')

doc.add_page_break()

# ============ 四、项目实现 ============
add_heading_cn(doc, '四、项目实现', 1)

add_body(doc, '4.0 项目目录结构', bold=True)
add_code(doc, '''HarmonyCockpit-backend/
├── config/
│   └── db.js                 # MySQL 连接池配置
├── controllers/
│   ├── usersController.js    # 用户模块控制器
│   ├── infosController.js    # 车辆信息控制器
│   ├── statesController.js   # 车辆控制控制器
│   └── feedbacksController.js# 意见反馈控制器
├── routes/
│   ├── usersRoutes.js        # 登录/用户路由
│   ├── infosRoutes.js        # 车辆信息路由
│   ├── statesRoutes.js       # 车辆控制路由
│   └── feedbacksRoutes.js    # 意见反馈路由
├── services/
│   ├── usersService.js       # 用户业务逻辑
│   ├── infosService.js       # 车辆信息业务逻辑
│   ├── statesService.js      # 车辆控制业务逻辑
│   └── feedbacksService.js   # 意见反馈业务逻辑
├── middlewares/
│   └── errorHandler.js       # 统一错误处理中间件
├── frontend/
│   ├── login.html            # 登录页
│   ├── index.html            # 主应用（SPA：首页/控制/统计/我的）
│   ├── style.css             # 移动端样式
│   ├── utils.js              # API工具函数
│   └── images/               # 车辆图片 + HIMA图标
├── app.js                    # Express 应用配置
├── server.js                 # 服务启动入口
├── init_db.sql               # 数据库初始化脚本
├── .env                      # 环境变量（端口/数据库配置）
└── package.json              # 依赖配置''')

add_img_placeholder(doc, '项目目录截图（1分）')

add_heading_cn(doc, '4.1 功能实现', 2)

# 4.1.1 登录模块
add_body(doc, '4.1.1 登录模块', bold=True)
add_body(doc, '用户通过用户ID和密码登录。后端采用 MVC 三层架构：Routes 处理路由分发，Controller 处理请求校验与响应，Service 层执行数据库查询。', indent=True)

add_body(doc, '核心代码 (services/usersService.js)：', indent=True)
add_code(doc, '''async function login(user_id, password) {
    const [rows] = await db.execute(
        'SELECT * FROM user WHERE user_id = ?', [user_id]);
    if (rows.length === 0) {
        const err = new Error('用户不存在');
        err.status = 401; throw err;
    }
    const user = rows[0];
    if (user.password !== password) {
        const err = new Error('密码错误');
        err.status = 401; throw err;
    }
    return { user_id, userName, Sex, Birthdate };
}''')

add_body(doc, '测试场景：', indent=True)
for s in ['用户登录成功（200 OK，返回用户信息）', '用户不存在（401，返回"用户不存在"）', '密码错误（401，返回"密码错误"）', '参数缺失（400，返回"参数不能为空"）']:
    add_body(doc, f'  · {s}', indent=True)

add_img_placeholder(doc, '登录模块 Postman 测试截图（成功+失败）')

# 4.1.2 查看/维护车主信息
add_body(doc, '4.1.2 查看/维护车主信息模块', bold=True)
add_body(doc, '支持查看和修改车主个人信息。通过 GET /api/users/:user_id 查询详情，PUT /api/users/:user_id 更新信息。更新采用动态字段更新策略，只更新客户端传递的字段。', indent=True)

add_code(doc, '''// 动态字段更新
async function update(user_id, data) {
    const fields = [], values = [];
    for (const [key, value] of Object.entries(data)) {
        if (value !== undefined && key !== 'user_id') {
            fields.push(key + ' = ?'); values.push(value);
        }
    }
    values.push(user_id);
    await db.execute('UPDATE user SET '+fields.join(',')+' WHERE user_id=?', values);
}''')

add_body(doc, '数据隔离：所有车辆查询通过 user_id 过滤，车主只能查看自己名下的车辆信息。', indent=True)

add_img_placeholder(doc, '车主信息模块 Postman 测试截图（查询详情+更新信息）')

# 4.1.3 车辆控制模块
add_body(doc, '4.1.3 车辆控制模块', bold=True)
add_body(doc, '核心功能模块，支持空调控制、车门灯光控制、车辆移动控制。所有操作通过 POST /api/states 创建新状态记录，每条记录携带完整 fields，确保各控制字段之间相互独立（如除雾与空调互不影响）。', indent=True)

add_body(doc, '空调控制状态机：', indent=True)
add_code(doc, '''const cur = state || {};
const nv = cur[key] === '1' ? '0' : '1';
// 独立发送：只改当前字段，保留其他字段不变
sendState({ [key]: nv });''')

add_body(doc, '车辆移动采用D-Pad设计：上下左右方向键为瞬时控制（按下触发+500ms自动恢复），中间停车按钮为锁定控制（按下保持、再次按下取消），停车激活时方向键自动锁止（灰色不可交互）。', indent=True)

add_code(doc, '''function doMove(key, val) {
    const cur = state || {};
    if (cur.stop1 === '1') return; // 停车锁定
    sendState({ [key]: val });
    if (val === '1') {
        setTimeout(() => sendState({ [key]: '0' }), 500); // 自动恢复
    }
}''')

add_body(doc, '操作记录：系统通过对比相邻状态记录的字段变化，自动生成操作历史（如"打开A/C"、"切换内循环"、"左转"、"停止前进"等），格式化展示为"用户XX控制车辆YY进行了ZZ操作"。', indent=True)

add_img_placeholder(doc, '车辆控制模块 Postman 测试截图（空调操作+车门操作+移动操作）')

# 4.1.4 意见反馈模块
add_body(doc, '4.1.4 意见反馈模块', bold=True)
add_body(doc, '支持反馈的完整 CRUD 操作，通过 LEFT JOIN user 表获取提交者姓名。反馈模块位于"我的"页面的二级菜单中，点击进入反馈列表，右下角"+"按钮新增反馈。', indent=True)

add_code(doc, '''async function list({ page, pageSize, user_id } = {}) {
    let sql = 'SELECT f.*, u.userName FROM feedback f
        LEFT JOIN user u ON f.user_id = u.user_id WHERE 1=1';
    // 动态拼接查询条件
    sql += ' ORDER BY f.time DESC LIMIT ? OFFSET ?';
    const [rows] = await db.execute(sql, params);
    return rows;
}''')

add_img_placeholder(doc, '意见反馈模块 Postman 测试截图（反馈列表+新增反馈）')

# 4.1.5 数据统计
add_body(doc, '4.1.5 数据统计模块', bold=True)
add_body(doc, '展示近7日行驶里程的 SVG 折线图（带渐变填充、数据点标注、网格线）。后端使用内存缓存，服务运行期间同用户同车辆的数据固定，重启服务后重新随机生成。', indent=True)

add_img_placeholder(doc, '数据统计页面截图')

# 提交 Git
add_body(doc, '4.2 提交 Git（1分）', bold=True)
add_body(doc, '项目代码已提交至 Git 仓库。', indent=True)
add_img_placeholder(doc, 'Git 提交截图')

doc.add_page_break()

# ============ 五、亮点功能 ============
add_heading_cn(doc, '五、亮点功能', 1)

add_body(doc, '除基本 CRUD 功能外，本系统还实现了以下亮点功能：', indent=True)

highlights = [
    ('1. 后端参数校验与统一错误处理',
     '所有接口均对请求参数进行非空校验和格式校验。通过 Express 中间件 errorHandler 统一捕获异常，返回标准化 JSON 错误响应（success: false, error: {code, message}），涵盖 400/401/404/500 等状态码。'),
    ('2. 操作行为追踪系统',
     '每次车辆控制操作（空调、车门、灯光、移动等）均 POST 一条完整状态记录到数据库。操作记录页面通过对比相邻记录的字段变化自动还原操作行为，格式化为"用户XX控制车辆YY进行了ZZ操作"，支持最多显示30条近期记录。'),
    ('3. D-Pad 车辆移动控制',
     '采用十字方向键（D-Pad）设计替代普通按钮，支持按下触发+自动恢复的瞬时控制模式。中间停车按钮为锁定开关，激活时所有方向键自动禁用（灰色不可交互），实现真实车辆控制逻辑。'),
    ('4. 动态字段更新（Dynamic Update）',
     '车辆信息更新接口支持按需传递字段，后端自动构建 UPDATE SQL，仅修改客户端指定的字段，避免全量覆写。'),
    ('5. 页面间状态实时同步',
     '首页与控制页共享同一 state 对象。每次操作后同时更新两页的 UI 元素，切换页面时直接使用本地状态刷新而非从服务器重新拉取，确保温度、开关状态等数据在两页间始终保持一致。'),
    ('6. SVG 数据可视化',
     '数据统计页使用纯 SVG 绘制折线图，包含渐变填充、数据点标注、网格线和日期标签，无需引入任何第三方图表库。'),
]
for title, desc in highlights:
    add_body(doc, title, bold=True, indent=True)
    add_body(doc, desc, indent=True)

add_img_placeholder(doc, '亮点功能截图（操作记录页、D-Pad控制、数据统计图等）')

doc.add_page_break()

# ============ 六、项目总结 ============
add_heading_cn(doc, '六、项目总结', 1)

add_body(doc, '1. 遇到的问题与解决方案', bold=True)

problems = [
    ('MySQL only_full_group_by 模式导致 GROUP BY 查询报错',
     '问题：统计里程的 SQL 中 SELECT 包含了非聚合列 i.miles，与 sql_mode=only_full_group_by 冲突。解决：使用 MAX(i.miles) 聚合函数替代直接查询，并修改 ORDER BY 使用 GROUP BY 中的 date 别名。'),
    ('前端按钮点击无效（this 指向问题）',
     '问题：toggle 按钮内嵌 <span> 图标，点击图标时 click 事件的 this 指向 span 而非父 div，导致 data-key 读取为 undefined。解决：使用 this.closest(\'.toggle-btn[data-key]\') 向上查找父按钮元素。'),
    ('状态记录重复主键冲突',
     '问题：快速操作时 PostgreSQL 的 state 表主键 (user_ID, car_ID, time) 中 time 仅精确到秒，同一秒内多次操作导致 Duplicate entry。解决：使用 ALTER TABLE 将 time 列改为 DATETIME(3) 毫秒精度，前端 nowStr() 函数生成包含毫秒的时间戳；同时添加 _sending 乐观锁防止并发 POST 重复提交。'),
    ('页面间状态不同步',
     '问题：首页和控制页分别从服务器拉取最新状态，可能拉回不同记录导致温度等数据显示不一致。解决：改为共享同一个 state 对象，操作后直接更新本地 state 并同时刷新两页 UI，页面切换时仅首次加载从服务器拉取，后续使用本地 state。'),
    ('异步刷新覆盖用户操作',
     '问题：POST 新状态后异步拉取服务器最新数据，回调中的旧数据覆盖了刚更新的本地状态。解决：移除 POST 后的异步刷新，仅在切换页面时首次加载从服务器获取，操作后依靠本地 state 保持一致性。'),
]
for q, a in problems:
    add_body(doc, f'· {q}', indent=True)
    add_body(doc, f'  解决：{a}', indent=True)

add_body(doc, '2. 项目收获', bold=True)
for item in [
    '掌握了 Node.js + Express + MySQL 全栈开发流程，深入理解了 MVC 分层架构中 Routes→Controller→Service→DB 的职责分离与数据流向。',
    '深入理解了 RESTful API 设计规范，包括 HTTP 方法语义（GET/POST/PUT/DELETE）、状态码使用、统一 JSON 响应格式、分页参数设计。',
    '通过实现数据隔离逻辑（user_id 过滤），理解了多用户系统中数据权限与安全的重要性。',
    '学会了 LEFT JOIN 多表关联查询与 SQL 聚合函数（MAX、COUNT、GROUP BY），以及视图和索引的创建与使用。',
    '实践了前端复杂交互开发：D-Pad 方向控制、停车锁定机制、页面间状态同步、SVG 图表绘制等。',
    '掌握了使用 rembg 库进行图片背景去除、PIL/Pillow 圆形裁剪等图像处理技能。',
]:
    add_body(doc, f'{item}', indent=True)

add_body(doc, '3. 不足之处', bold=True)
for item in [
    '密码采用明文存储，后续应引入 bcrypt 加密和 JWT Token 身份认证机制以提升安全性。',
    '前端未实现完整的输入参数白名单校验，存在潜在的安全风险，后续可加入前端表单验证。',
    '未实现日志记录和请求追踪功能，后续可引入 Winston 等日志库。',
    '车辆控制操作目前全部通过 POST 创建新记录，数据量较大时可考虑定期归档历史记录。',
]:
    add_body(doc, f'{item}', indent=True)

# ============ 保存 ============
output_path = 'C:/Users/15836/Documents/孙延煦/编程实践/HarmonyCockpit-backend/四级项目报告.docx'
doc.save(output_path)
print(f'报告已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')