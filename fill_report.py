from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document('05 形成性考核-四级项目-项目报告模板.docx')

def fill_para(para, new_text):
    if not para.runs:
        run = para.add_run(new_text)
        run.font.name = '宋体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(12)
        return
    for run in para.runs: run.text = ''
    para.runs[0].text = new_text

def insert_table_after(anchor_para, table):
    """Insert table right after anchor paragraph in document"""
    anchor_para._element.addnext(table._tbl)

def insert_para_after(anchor_para, new_para):
    anchor_para._element.addnext(new_para._element)

def make_api_table_after(doc, anchor, desc, url, method, req_params, res_type, res_example):
    """Create API table and insert after anchor paragraph"""
    table = doc.add_table(rows=6, cols=2, style='Table Grid')

    # Move table from end to right after anchor
    anchor._element.addnext(table._tbl)

    rows_data = [
        ('功能说明', desc),
        ('URL地址', url),
        ('请求方式', method),
    ]

    # Fill first 3 rows
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Row 3: 请求说明 with parameter details
    table.rows[3].cells[0].text = '请求说明'
    if req_params:
        lines = ['参数名\t是否必填\t类型\t说明']
        for p in req_params:
            lines.append(f'{p[0]}\t{"是" if p[1] else "否"}\t{p[2]}\t{p[3]}')
        table.rows[3].cells[1].text = '\n'.join(lines)
    else:
        table.rows[3].cells[1].text = '无'

    # Row 4: 返回参数
    table.rows[4].cells[0].text = '返回参数'
    table.rows[4].cells[1].text = res_type

    # Row 5: 返回示例
    table.rows[5].cells[0].text = '返回示例'
    table.rows[5].cells[1].text = res_example

    # Add spacing paragraph after table
    spacer = doc.add_paragraph()
    anchor._element.addnext(spacer._element)

    return table

paras = doc.paragraphs

# ============================================================
# Section 1: 需求分析 [20-26]
# ============================================================
fill_para(paras[21], '本项目为鸿蒙智能座舱车主端后端系统，采用 Node.js + Express + MySQL 技术栈，MVC 分层架构，为车主端提供 RESTful API 服务。')
fill_para(paras[22], '1. 登录认证：车主通过用户ID和密码登录，验证后返回个人信息。支持用户不存在、密码错误等异常处理。')
fill_para(paras[23], '2. 远程车辆控制：支持空调开关及温度调节（16~32°C）、内外循环、除雾、车门开关、灯光开关、车辆移动（前进/后退/左转/右转/停车）。所有操作记录到state表，可追溯。')
fill_para(paras[24], '3. 提交反馈或建议：提交使用反馈，分页查看所有反馈列表。支持新增/更新/删除。')
fill_para(paras[25], '4. 查看和维护个人信息：查看姓名、性别、出生日期，支持修改（动态字段更新）。')
fill_para(paras[26], '5. 近7日车辆行驶里程统计：SVG折线图展示，含总里程、日均里程、最长单日统计。')

# ============================================================
# Section 2: 数据库设计 [28-31]
# ============================================================
fill_para(paras[29], '鸿蒙智能座舱车主端数据库 car，MySQL 8.0，utf8mb4，4张核心表+2视图+2索引。')

# Replace table 0 (sys_user -> user table)
t0 = doc.tables[1]
while len(t0.rows) > 1:
    t0.rows[1]._tr.getparent().remove(t0.rows[1]._tr)
for i, h in enumerate(['字段名', '类型', '约束', '说明']):
    t0.rows[0].cells[i].text = h
for r in [['user_id','CHAR(15)','PRIMARY KEY','用户ID'],['userName','VARCHAR(20)','NOT NULL','姓名'],['Sex','VARCHAR(2)','NOT NULL','性别'],['Birthdate','DATE','NOT NULL','出生日期'],['password','VARCHAR(20)','NOT NULL','密码']]:
    row = t0.add_row()
    for i, v in enumerate(r): row.cells[i].text = v

fill_para(paras[30], '表1 user表：存储车主用户信息。info表：存储车辆信息（car_id,carname,brand,frame_id,user_id,power,max_power,torque,battery,miles,mpg）。state表：存储车辆操控状态，复合主键(user_ID,car_ID,time)，time为DATETIME(3)毫秒精度。feedback表：存储意见反馈（idea_id,ideaName,content,time,user_id）。视图：xuqing_car、xuqing_carstate。索引：IDX_mpg、IDX_miles。')
fill_para(paras[31], '')

# ============================================================
# Section 3: 接口设计 [42-52] - DETAILED FORMAT
# ============================================================
fill_para(paras[43], '接口设计详细如下（Base URL: http://localhost:3000/api）：')

# Remove old example table from document body
t2 = doc.tables[2]
body = doc.element.body
body.remove(t2._tbl)

# Clear old list paragraphs
for i in [44,45,46,47,48,49,50,51,52]:
    if i < len(paras):
        fill_para(paras[i], '')

# Create a spacing paragraph after the heading
spacer = doc.add_paragraph()
paras[43]._element.addnext(spacer._element)
current_anchor = spacer

# --- API 1: 用户登录 ---
make_api_table_after(doc, current_anchor, '用户登录', 'http://localhost:3000/api/users/login', 'POST',
    [['user_id','是','string','用户ID'],['password','是','string','密码']],
    'JSON',
    '{"success":true,"data":{"user_id":"277273737937","userName":"王琦","Sex":"男","Birthdate":"2009-03-01"},"message":"登录成功"}')

# --- API 2: 查询用户详情 ---
s2 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s2._element)
make_api_table_after(doc, s2, '查询用户详情', 'http://localhost:3000/api/users/{user_id}', 'GET',
    [['user_id','是','string','用户ID(path)']],
    'JSON',
    '{"success":true,"data":{"user_id":"277273737937","userName":"王琦","Sex":"男","Birthdate":"2009-03-01"}}')

# --- API 3: 更新用户信息 ---
s3 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s3._element)
make_api_table_after(doc, s3, '更新用户信息', 'http://localhost:3000/api/users/{user_id}', 'PUT',
    [['userName','否','string','姓名'],['Sex','否','string','性别'],['Birthdate','否','string','出生日期'],['password','否','string','密码']],
    'JSON',
    '{"success":true,"message":"更新成功"}')

# --- API 4: 查询车辆列表 ---
s4 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s4._element)
make_api_table_after(doc, s4, '分页查询车辆列表（带车主姓名）', 'http://localhost:3000/api/infos', 'GET',
    [['page','否','number','页码(默认1)'],['pageSize','否','number','每页条数(默认10)'],['user_id','否','string','按用户ID过滤']],
    'JSON',
    '{"success":true,"data":{"list":[...],"pagination":{"page":1,"pageSize":10,"total":11}}}')

# --- API 5: 新增/更新/删除车辆 ---
s5 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s5._element)
make_api_table_after(doc, s5, '新增/更新/删除车辆', 'http://localhost:3000/api/infos[/{car_id}]', 'POST/PUT/DELETE',
    [['car_id','是','string','车辆ID'],['carName','是','string','车辆名称'],['brand','是','string','品牌'],['frame_id','是','string','车架号'],['user_id','是','string','用户ID']],
    'JSON',
    '{"success":true,"message":"操作成功"}')

# --- API 6: 查询车辆状态 ---
s6 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s6._element)
make_api_table_after(doc, s6, '分页查询车辆操控状态', 'http://localhost:3000/api/states', 'GET',
    [['page','否','number','页码'],['pageSize','否','number','每页条数'],['user_ID','否','string','用户ID'],['car_ID','否','string','车辆ID']],
    'JSON',
    '{"success":true,"data":[{...}],"meta":{"total":31}}')

# --- API 7: 查询最新状态 ---
s7 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s7._element)
make_api_table_after(doc, s7, '查询车辆最新状态', 'http://localhost:3000/api/states/latest', 'GET',
    [['user_ID','是','string','用户ID(query)'],['car_ID','是','string','车辆ID(query)']],
    'JSON',
    '{"success":true,"data":{"ac_on":"0","ac_temp":22,"door":"0",...}}')

# --- API 8: 新增操控记录 ---
s8 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s8._element)
make_api_table_after(doc, s8, '新增车辆操控记录', 'http://localhost:3000/api/states', 'POST',
    [['user_ID','是','string','用户ID'],['car_ID','是','string','车辆ID'],['ac_on','否','string','空调(0/1)'],['ac_temp','否','float','温度(℃)'],['door','否','string','车门(0/1)'],['light','否','string','灯光(0/1)'],['time','是','string','时间(DATETIME(3))']],
    'JSON',
    '{"success":true,"message":"新增成功"}')

# --- API 9: 里程统计 ---
s9 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s9._element)
make_api_table_after(doc, s9, '近7日行驶里程统计', 'http://localhost:3000/api/states/stats/miles', 'GET',
    [['user_ID','是','string','用户ID'],['car_ID','是','string','车辆ID']],
    'JSON',
    '{"success":true,"data":[{"date":"2026-06-30","miles":31},...]}')

# --- API 10: 分页查询反馈 ---
s10 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s10._element)
make_api_table_after(doc, s10, '分页查询意见反馈', 'http://localhost:3000/api/feedbacks', 'GET',
    [['page','否','number','页码'],['pageSize','否','number','每页条数'],['user_id','否','string','按用户过滤']],
    'JSON',
    '{"success":true,"data":{"list":[{...}]}}')

# --- API 11: 新增反馈 ---
s11 = doc.add_paragraph()
doc.paragraphs[-2]._element.addnext(s11._element)
make_api_table_after(doc, s11, '新增意见反馈', 'http://localhost:3000/api/feedbacks', 'POST',
    [['idea_id','是','string','反馈ID'],['ideaName','是','string','反馈标题'],['content','是','string','反馈内容'],['user_id','是','string','用户ID']],
    'JSON',
    '{"success":true,"message":"提交成功"}')

# ============================================================
# Section 4: 项目实现 [54-75]
# ============================================================
fill_para(paras[55], '项目采用 MVC 分层架构，目录结构如下：')

# Insert code block after para[55]
code_p = doc.add_paragraph()
paras[55]._element.addnext(code_p._element)
run = code_p.add_run(
'HarmonyCockpit-backend/\n'
'├── config/db.js .env             # 配置层(数据库/环境变量)\n'
'├── routes/       # 路由层(4个路由文件)\n'
'├── controllers/  # 控制层(参数校验→调用service)\n'
'├── services/     # 业务层(数据库CRUD+LEFT JOIN)\n'
'├── middlewares/  # 中间件(errorHandler统一错误处理)\n'
'├── frontend/     # 移动端SPA前端\n'
'│   ├── login.html index.html style.css utils.js\n'
'│   └── images/ (M7/S7/S800车图+HIMA图标)\n'
'├── init_db.sql app.js server.js package.json'
)
run.font.name = 'Consolas'; run.font.size = Pt(9)

fill_para(paras[56], '【📷 此处插入项目目录截图（1分）】')
fill_para(paras[59], '4.1 功能实现（5分）：车主端实现核心功能，严格数据隔离和列表显示。')

# 登录模块
fill_para(paras[61], '1. 登录模块：Routes→Controller→Service→DB四层调用链，支持成功/失败/异常全场景。')
fill_para(paras[62], '''核心代码 (services/usersService.js):
async function login(user_id, password) {
    const [rows] = await db.execute('SELECT * FROM user WHERE user_id=?', [user_id]);
    if (rows.length === 0) throw {status:401, message:'用户不存在'};
    if (rows[0].password !== password) throw {status:401, message:'密码错误'};
    return { user_id, userName, Sex, Birthdate };
}''')
fill_para(paras[63], '测试场景：登录成功(200+用户信息)、用户不存在(401)、密码错误(401)、参数缺失(400)')
fill_para(paras[64], '【📷 此处插入登录成功 + 失败 Postman截图】')

# 控制模块
fill_para(paras[67], '2. 车辆状态控制：POST /api/states 创建新记录，每条记录含14个控制字段。')
fill_para(paras[68], '''核心逻辑：每次点击POST完整状态快照，操作记录页对比相邻记录自动还原操作行为。
例："用户王琦控制车辆M7进行了打开A/C操作" — 2026-07-06 14:45:00''')
fill_para(paras[70], '【📷 此处插入车辆控制Postman截图（空调+车门+移动）】')
fill_para(paras[72], '3. 意见反馈：完整CRUD + LEFT JOIN user获取提交者姓名，位于"我的"页二级菜单。')
fill_para(paras[74], '【📷 此处插入反馈模块Postman截图】')

fill_para(paras[75], '4. 提交Git（1分）：【📷 此处插入 Git 提交截图】')

# ============================================================
# Section 5: 亮点功能 [76-81]
# ============================================================
fill_para(paras[77], '1. 操作行为追踪系统：每次控制操作POST完整状态记录，操作记录页对比相邻记录字段变化自动生成操作历史（"打开A/C""左转""停止前进"等），最多显示30条。')
fill_para(paras[78], '2. D-Pad车辆移动控制：十字方向键瞬时触发（按下即发，500ms自动恢复），中间停车按钮锁定控制（激活时方向键自动禁用变灰）。')
fill_para(paras[79], '3. 动态字段更新：info和state更新接口支持按需传递字段，后端自动构建UPDATE SQL，避免全量覆写。')
fill_para(paras[81], '4. 页面间状态实时同步：首页与控制页共享state对象，操作后同时刷新两页UI；切换页面使用本地状态，无需重新拉取服务器数据。')
# [82] might be empty
fill_para(paras[85], '5. SVG数据可视化：统计页使用纯SVG绘制折线图（梯度填充+数据点+网格线），无需第三方图表库。')
fill_para(paras[87], '6. 错误处理与状态码规范：统一errorHandler中间件+参数校验，覆盖400/401/404/500。')
fill_para(paras[88], '【📷 此处插入亮点功能截图（操作记录页、D-Pad控制、统计图）】')

# ============================================================
# Section 6: 项目总结 [84-92]
# ============================================================
fill_para(paras[89], '''1 遇到的问题：
问题1: MySQL only_full_group_by导致统计查询报错 → 使用MAX(i.miles)聚合函数
问题2: 前端按钮点击this指向span → 改用closest('.toggle-btn[data-key]')
问题3: state表主键冲突(Duplicate entry) → time改为DATETIME(3)毫秒精度+前端乐观锁
问题4: 页面间状态不同步 → 共享state对象+两页UI同步刷新
问题5: 异步刷新覆盖用户操作 → 移除POST后异步刷新，首载从服务器拉取''')
fill_para(paras[90], '''2 项目收获：
1. Node.js+Express+MySQL全栈开发，MVC架构理解
2. RESTful API设计：HTTP方法语义、状态码、统一响应
3. 数据隔离(user_id过滤)与LEFT JOIN关联查询
4. SQL聚合函数、视图、索引、DATETIME(3)精度处理
5. 前端：D-Pad控制、状态同步、SVG绘制、rembg图像处理''')
fill_para(paras[92], '''不足：密码明文存储(后续用bcrypt+JWT)、未实现日志(Winston)、无表单前端验证、无鸿蒙原生UI''')

# Save
output_path = '四级项目报告_已填写.docx'
doc.save(output_path)
print(f'Saved: {output_path}')